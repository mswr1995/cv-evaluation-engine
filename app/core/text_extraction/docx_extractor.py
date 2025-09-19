"""Text extractor for DOCX files."""

import logging
from pathlib import Path
from typing import Tuple

try:
    from docx import Document
except ImportError:
    Document = None

from app.core.text_extraction.base import TextExtractor

logger = logging.getLogger(__name__)


class DocxExtractor(TextExtractor):
    """Text extractor for DOCX files (.docx)."""
    
    def __init__(self):
        """Initialize DOCX extractor and check dependencies."""
        if Document is None:
            logger.error("python-docx library not installed. DOCX extraction will not work.")
    
    def extract_text(self, file_path: str) -> Tuple[bool, str, str]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple[bool, str, str]: (success, extracted_text, error_message)
        """
        if Document is None:
            error_msg = "python-docx library not installed"
            logger.error(error_msg)
            return False, "", error_msg
        
        logger.debug(f"Extracting text from DOCX file: {file_path}")
        
        try:
            # Verify file exists
            if not Path(file_path).exists():
                error_msg = f"File not found: {file_path}"
                logger.error(error_msg)
                return False, "", error_msg
            
            # Read DOCX file
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            extracted_paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    extracted_paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        extracted_paragraphs.append(" | ".join(row_text))
            
            # Combine all extracted text
            full_text = "\n".join(extracted_paragraphs)
            
            if not full_text.strip():
                error_msg = "No readable text found in DOCX file"
                logger.warning(error_msg)
                return False, "", error_msg
            
            logger.debug(f"Successfully extracted {len(full_text)} characters from DOCX")
            return True, full_text, ""
            
        except Exception as e:
            error_msg = f"Failed to extract text from DOCX: {str(e)}"
            logger.error(error_msg)
            return False, "", error_msg
    
    def supports_file_type(self, file_extension: str) -> bool:
        """Check if this extractor supports DOCX files."""
        return file_extension.lower() == "docx"